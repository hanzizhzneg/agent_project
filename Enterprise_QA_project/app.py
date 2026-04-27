from __future__ import annotations

import argparse
import os
import json
from pathlib import Path
from typing import Iterator, Literal

from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field
import pymupdf as fitz
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI


StoreType = Literal["chroma", "faiss"]
BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "knowledge_base"
CHROMA_DIR = BASE_DIR / "vector_db" / "chroma"
# FAISS on Windows may fail with non-ASCII paths; default to a user-home ASCII-friendly path.
FAISS_DIR = Path(
    os.getenv(
        "FAISS_DIR",
        str(Path.home() / "enterprise_qa_data" / "vector_db" / "faiss"),
    )
)
WEB_DIR = BASE_DIR / "web"
INDEX_HTML = WEB_DIR / "index.html"
TEST_HTML = WEB_DIR / "test.html"


class IngestRequest(BaseModel):
    docs_dir: str = Field(default=str(DATA_DIR), description="PDF/Word docs directory")
    store_type: StoreType = "chroma"
    chunk_size: int = 800
    chunk_overlap: int = 150


class AskRequest(BaseModel):
    question: str = Field(..., min_length=1)
    store_type: StoreType = "chroma"
    top_k: int = 4
    chat_history: list[dict[str, str]] = Field(
        default_factory=list,
        description="Optional history. Example: [{'role':'user','content':'...'}]",
    )


class AskResponse(BaseModel):
    answer: str
    contexts: list[str]
    sources: list[str]


def get_embeddings() -> HuggingFaceEmbeddings:
    # First run downloads a multilingual sentence-transformers model.
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )


def get_llm() -> ChatOpenAI:
    api_key = os.getenv("LLM_API_KEY")
    model_name = os.getenv("LLM_MODEL", "qwen3.6-35b-a3b")
    base_url = os.getenv(
        "LLM_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1"
    )
    temperature = float(os.getenv("LLM_TEMPERATURE", "0.2"))

    if not api_key:
        raise ValueError("Missing environment variable: LLM_API_KEY (DashScope/Bailian key)")

    return ChatOpenAI(
        model=model_name,
        api_key=api_key,
        base_url=base_url,
        temperature=temperature,
    )


def load_pdf(file_path: Path) -> list[Document]:
    pages: list[Document] = []
    # Prefer PyMuPDF for better Chinese PDF extraction quality.
    try:
        with fitz.open(str(file_path)) as pdf:
            for idx, page in enumerate(pdf, start=1):
                text = page.get_text("text") or ""
                if text.strip():
                    pages.append(
                        Document(
                            page_content=text,
                            metadata={"source": str(file_path), "page": idx, "type": "pdf"},
                        )
                    )
        if pages:
            return pages
    except Exception:
        # Fall back to PyPDF2 when PyMuPDF fails on specific PDFs.
        pages = []

    reader = PdfReader(str(file_path))
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(
                Document(
                    page_content=text,
                    metadata={"source": str(file_path), "page": idx, "type": "pdf"},
                )
            )
    return pages


def load_docx(file_path: Path) -> list[Document]:
    doc = DocxDocument(str(file_path))
    content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not content.strip():
        return []
    return [
        Document(
            page_content=content,
            metadata={"source": str(file_path), "type": "docx"},
        )
    ]


def load_documents(docs_dir: Path) -> list[Document]:
    if not docs_dir.exists() or not docs_dir.is_dir():
        raise ValueError(f"Document directory not found: {docs_dir}")

    docs: list[Document] = []
    for file_path in docs_dir.rglob("*"):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() == ".pdf":
            docs.extend(load_pdf(file_path))
        elif file_path.suffix.lower() in {".docx"}:
            docs.extend(load_docx(file_path))

    if not docs:
        raise ValueError(f"No PDF/DOCX files found in: {docs_dir}")
    return docs


def split_documents(
    docs: list[Document], chunk_size: int, chunk_overlap: int
) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    return splitter.split_documents(docs)


def build_vector_store(
    docs: list[Document], store_type: StoreType, chunk_size: int, chunk_overlap: int
) -> None:
    embeddings = get_embeddings()
    split_docs = split_documents(docs, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    if store_type == "chroma":
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
        vectorstore = Chroma.from_documents(
            documents=split_docs,
            embedding=embeddings,
            persist_directory=str(CHROMA_DIR),
            collection_name="enterprise_qa",
        )
        # langchain-chroma newer versions persist automatically and no longer
        # expose `persist()`. Keep backward compatibility with older versions.
        if hasattr(vectorstore, "persist"):
            vectorstore.persist()
        return

    FAISS_DIR.mkdir(parents=True, exist_ok=True)
    vectorstore = FAISS.from_documents(split_docs, embeddings)
    vectorstore.save_local(str(FAISS_DIR))


def load_vector_store(store_type: StoreType):
    embeddings = get_embeddings()
    if store_type == "chroma":
        if not CHROMA_DIR.exists():
            raise ValueError("Chroma vector store not found, please ingest first.")
        return Chroma(
            persist_directory=str(CHROMA_DIR),
            embedding_function=embeddings,
            collection_name="enterprise_qa",
        )

    if not FAISS_DIR.exists():
        raise ValueError(
            f"FAISS vector store not found at {FAISS_DIR}, please ingest first."
        )
    return FAISS.load_local(
        str(FAISS_DIR),
        embeddings,
        allow_dangerous_deserialization=True,
    )


def format_source(doc: Document) -> str:
    source = doc.metadata.get("source", "unknown")
    page = doc.metadata.get("page")
    if page is None:
        return str(source)
    return f"{source} (page {page})"


def build_history_text(chat_history: list[dict[str, str]]) -> str:
    lines: list[str] = []
    for msg in chat_history:
        role = msg.get("role", "").strip().lower()
        content = msg.get("content", "").strip()
        if role in {"user", "assistant"} and content:
            label = "用户" if role == "user" else "助手"
            lines.append(f"{label}: {content}")
    return "\n".join(lines)


def retrieve_context(
    question: str, store_type: StoreType, top_k: int
) -> tuple[str, list[str], list[str]]:
    vectorstore = load_vector_store(store_type)
    results = vectorstore.similarity_search(question, k=top_k)
    if not results:
        return "", [], []

    contexts = [doc.page_content for doc in results]
    sources = [format_source(doc) for doc in results]
    context_text = "\n\n".join(
        f"[片段{i + 1}]\n{content}" for i, content in enumerate(contexts)
    )
    return context_text, contexts, sources


def build_rag_answer(
    question: str,
    context_text: str,
    chat_history: list[dict[str, str]],
) -> str:
    history_text = build_history_text(chat_history)
    history_block = history_text if history_text else "无历史对话"

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "你是企业知识库问答助手。请只依据提供的上下文回答。"
                    "如果上下文无法支持答案，请明确说“根据当前知识库无法确定”。"
                    "回答尽量简洁，并在末尾附上你引用的片段编号。"
                ),
            ),
            (
                "human",
                "历史对话：\n{history_block}\n\n问题：{question}\n\n可用上下文：\n{context_text}",
            ),
        ]
    )

    llm = get_llm()
    chain = prompt | llm
    result = chain.invoke(
        {
            "question": question,
            "context_text": context_text,
            "history_block": history_block,
        }
    )
    return result.content if isinstance(result.content, str) else str(result.content)


def answer_question(
    question: str,
    store_type: StoreType,
    top_k: int = 4,
    chat_history: list[dict[str, str]] | None = None,
) -> AskResponse:
    history = chat_history or []
    context_text, contexts, sources = retrieve_context(question, store_type, top_k)
    if not contexts:
        return AskResponse(answer="未检索到相关内容。", contexts=[], sources=[])

    answer = build_rag_answer(
        question=question,
        context_text=context_text,
        chat_history=history,
    )
    return AskResponse(answer=answer, contexts=contexts, sources=sources)


def stream_answer_chunks(
    question: str,
    store_type: StoreType,
    top_k: int = 4,
    chat_history: list[dict[str, str]] | None = None,
) -> Iterator[str]:
    history = chat_history or []
    context_text, contexts, sources = retrieve_context(question, store_type, top_k)
    if not contexts:
        payload = {"type": "done", "answer": "未检索到相关内容。", "sources": []}
        yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"
        return

    history_text = build_history_text(history)
    history_block = history_text if history_text else "无历史对话"
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "你是企业知识库问答助手。请只依据提供的上下文回答。"
                    "如果上下文无法支持答案，请明确说“根据当前知识库无法确定”。"
                    "回答尽量简洁，并在末尾附上你引用的片段编号。"
                ),
            ),
            (
                "human",
                "历史对话：\n{history_block}\n\n问题：{question}\n\n可用上下文：\n{context_text}",
            ),
        ]
    )

    llm = get_llm()
    chain = prompt | llm
    full_text = ""
    for chunk in chain.stream(
        {
            "question": question,
            "context_text": context_text,
            "history_block": history_block,
        }
    ):
        text = chunk.content if isinstance(chunk.content, str) else str(chunk.content)
        if not text:
            continue
        full_text += text
        yield f"data: {json.dumps({'type': 'token', 'text': text}, ensure_ascii=False)}\n\n"

    final_payload = {
        "type": "done",
        "answer": full_text,
        "sources": sources,
        "contexts": contexts,
    }
    yield f"data: {json.dumps(final_payload, ensure_ascii=False)}\n\n"


app = FastAPI(title="LangChain PDF/Word QA Demo")


@app.get("/")
def index():
    if not INDEX_HTML.exists():
        raise HTTPException(status_code=404, detail="Frontend page not found.")
    return FileResponse(str(INDEX_HTML))


@app.get("/test")
def test_page():
    if not TEST_HTML.exists():
        raise HTTPException(status_code=404, detail="Test frontend page not found.")
    return FileResponse(str(TEST_HTML))


@app.post("/ingest")
def ingest(req: IngestRequest):
    try:
        docs = load_documents(Path(req.docs_dir))
        build_vector_store(
            docs=docs,
            store_type=req.store_type,
            chunk_size=req.chunk_size,
            chunk_overlap=req.chunk_overlap,
        )
        return {
            "ok": True,
            "message": "Ingestion completed.",
            "docs_count": len(docs),
            "store_type": req.store_type,
        }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Failed to ingest: {exc}") from exc


@app.post("/ask", response_model=AskResponse)
def ask(req: AskRequest):
    try:
        return answer_question(
            question=req.question,
            store_type=req.store_type,
            top_k=req.top_k,
            chat_history=req.chat_history,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Failed to answer: {exc}") from exc


@app.post("/ask_stream")
def ask_stream(req: AskRequest):
    try:
        generator = stream_answer_chunks(
            question=req.question,
            store_type=req.store_type,
            top_k=req.top_k,
            chat_history=req.chat_history,
        )
        return StreamingResponse(generator, media_type="text/event-stream")
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Failed to stream answer: {exc}") from exc


def main() -> None:
    parser = argparse.ArgumentParser(description="PDF/Word QA demo")
    subparsers = parser.add_subparsers(dest="command", required=True)

    ingest_parser = subparsers.add_parser("ingest", help="Ingest docs into vector store")
    ingest_parser.add_argument("--docs-dir", default=str(DATA_DIR))
    ingest_parser.add_argument("--store-type", choices=["chroma", "faiss"], default="chroma")
    ingest_parser.add_argument("--chunk-size", type=int, default=800)
    ingest_parser.add_argument("--chunk-overlap", type=int, default=150)

    ask_parser = subparsers.add_parser("ask", help="Ask a question")
    ask_parser.add_argument("--question", required=True)
    ask_parser.add_argument("--store-type", choices=["chroma", "faiss"], default="chroma")
    ask_parser.add_argument("--top-k", type=int, default=4)

    args = parser.parse_args()

    if args.command == "ingest":
        docs = load_documents(Path(args.docs_dir))
        build_vector_store(
            docs=docs,
            store_type=args.store_type,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
        )
        print(f"Ingest done. docs={len(docs)}, store={args.store_type}")
    elif args.command == "ask":
        result = answer_question(
            question=args.question, store_type=args.store_type, top_k=args.top_k
        )
        print(result.answer)


if __name__ == "__main__":
    main()
